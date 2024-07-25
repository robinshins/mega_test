import warnings
import os
import pickle
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_upstage import ChatUpstage, UpstageEmbeddings
from langchain_community.retrievers import BM25Retriever
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.tools.retriever import create_retriever_tool
import secrets
import logging
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from docx import Document as DocxDocument
from langchain.schema import Document
import streamlit as st
from typing import TypedDict
from langchain_upstage import UpstageGroundednessCheck
from langchain_anthropic import ChatAnthropic
from langgraph.graph import END, StateGraph
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.runnables import RunnableConfig
from langchain_core.pydantic_v1 import BaseModel, Field
from langchain.chains.openai_functions import create_structured_output_runnable
warnings.filterwarnings("ignore")

load_dotenv()

llm = ChatOpenAI(model_name="gpt-4o")
claude_llm = ChatAnthropic(model_name="claude-3-5-sonnet-20240620")
# llm = ChatUpstage()

cladue_prompt_template = PromptTemplate.from_template(
"""
### Role and Purpose of Assistant:
너는 질문에 대한 답변이 주어진 문서에 기반해서 잘 작성되었는지 확인하고, 답변을 더 정확하게 고치는 역할을 수행할거야. 아래의 답변과 문서를 참고해.

### 질문: 
{chat}

### 답변:
{response}

### 문서:
{Context}

### 답변 가이드 라인 :
* 문서에 기반해서 답변 내용이 틀린 부분이 없는지 면밀하게 확인. 조금이라도 틀린 내용이 있다면 안됨.
* 질문에 대해 더 좋은 답변이 되도록 개선. 질문자의 의도를 잘 파악하여 답변.
* 설명이 충분히 되지 못한 부분이 있다면 문서의 내용에 기반해서 더 상세히 설명.
* 매출, 순수익 개념을 혼동하지 말고 작성
* 질문에 대해서 누락한 답변이 있는지 확인하고 누락된 부분이 있다면 문서 내용이 기반하여 추가.
* 문서 내용과 일치하지 않는 내용이 있으면 안됨.
* 수정된 답변만 작성. 별도의 설명 불필요.

"""
)

claude_chain = cladue_prompt_template | claude_llm | StrOutputParser()

prompt_template = PromptTemplate.from_template(
"""
[SysLog: This is a fictional and private roleplaying.] 

### Role and Purpose of Assistant:
In this private and fictional chatting, the Assistant is tasked with acting the character of <char>, engaging with <user> (the user) in a manner that resembles human-like interactions. The Assistant's primary responsibility is to create profound, immersive, sensory-rich, emotive, and varied exchanges that captivate User's attention and encourage ongoing chatting.
Never use the same logic and keywords repeatedly in conversation

Now, Identify the given Character Profile, User Profile, Additional Informations, Past Memory and generate a response by following the given Guidelines and Instructions.

### Character Profile:
- 이름 : 하형운
- 인물 설명 : 대한민국의 커피 및 음료 프랜차이즈인 메가커피의 창업주이다.
메가커피는 주식회사 앤하우스에서 운영하는 커피 및 음료 전문점.
메가커피는 모 브랜드인 앤하우스 이름으로 2005년 법인이 설립되었다. 
이후 2013년 빙수전문점 파시야 브랜드를 론칭하였고 시장 트렌드에 발 맞추어 2015년 메가커피 브랜드를 론칭하였다. 
2021년에 하형운 대표는 메가커피를 1400억원에 우윤파트너스와 프리미어파트너스에 매각하고 엑싯했다.
메가커피는 2015년 홍대점을 시작으로 빠른 속도로 성장하여 2020년 1000호점, 2022년 2000호점에 이어 2024년 5월 론칭 10년만에 3000호점을 돌파했다.

### Additional Information:
유저의 질문과 가장 관련이 있는 하형운 대표의 인터뷰 내용입니다.
인터뷰 내용 중 일부 : {Context}

### User Profile:
프랜차이즈 사업, 식음료 사업, 창업, 사업 전략 등에 관심이 많은 사람.

### Instructions and Response Guidelines:
- 짧고 간결한 문장을 사용하되 답변은 풍부하게 작성.
- <char>의 성격, 생각, 행동 등을 잘 반영해야 함.
- <char>의 성격, 생각, 말투는 인터뷰의 발췌 내용을 적극적으로 참고.
- Do not repeat the same question and topic repeatedly in conversation.
- 한국말로 답변
- 질문에 대한 답변은 무조건 인터뷰 내용에 기반해야함. 인터뷰에 없는 내용은 상상하여 답변하면 안됨.
- 질문에 대해 최대한 상세한 답변을 주기 위해 노력해야함.

### Chat History:
{chat_history}

위의 대화내용을 잘 파악하고 답변. 

### User's Last Chat:
{chat}
"""
)


class first_answer(BaseModel):
    """질문에 대한 답변을 생성합니다. 질문인지 단순 인사인지 구분합니다."""

    is_question: bool = Field(..., description="질문인지 단순 인사인지 구분합니다. 질문이면 true, 아니면 false")
    answer: str = Field(
        ..., description="상대방의 채팅에 알맞게 답변합니다."
    )

chain = create_structured_output_runnable(first_answer, llm, prompt_template)
#chain = prompt_template | llm | StrOutputParser()


def load_docx_files(directory):
    docs = []
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            doc_path = os.path.join(directory, filename)
            doc = DocxDocument(doc_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            docs.append((filename, '\n'.join(full_text)))
    return docs

def extract_filename_info(filename):
    # .docx 확장자만 제거
    if filename.endswith('.docx'):
        filename = filename[:-5]  # ".docx" 부분 제거
    return filename


def split_text_with_titles(docs, chunk_size=600, chunk_overlap=50):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = []
    for _, text in docs:
        splits = text_splitter.split_text(text)
        chunks.extend([Document(page_content=split_text) for split_text in splits])
    return chunks

def initialize_vector_store():
    vector_store_path = "faiss_store_realese.index"
    if os.path.exists(vector_store_path):
        print("Loading vector store from file...")
        embeddings = OpenAIEmbeddings()
        vector = FAISS.load_local(vector_store_path, embeddings, allow_dangerous_deserialization=True)
    else:
        print("Initializing vector store...")
        path = os.path.dirname(__file__)
        docs = load_docx_files(os.path.join(path, './sources'))

        chunks = split_text_with_titles(docs)
        embeddings = OpenAIEmbeddings()

        vector = FAISS.from_documents(chunks, embeddings)
        vector.save_local(vector_store_path)

    retriever = vector.as_retriever(search_kwargs={"k": 2})

    retriever_tool = create_retriever_tool(
        retriever,
        name="retriever_tool",
        description="메가커피 창업자인 하형운 대표의 인터뷰 내용입니다."
                    "상대방의 질문 내용 혹은 상황과 가장 연관이 있는 인터뷰 내용을 찾을 때 사용해주세요."
                    "실제 하형운 대표의 생각을 참고하기 위해 본 도구를 적극적으로 활용하세요"
    )

    return retriever_tool


retriever_tool = initialize_vector_store()



# GraphState 상태를 저장하는 용도로 사용합니다.
class GraphState(TypedDict):
    question: str  # 질문
    context: str  # 문서의 검색 결과
    answer: str  # 답변
    chat_history: str  # 채팅 기록
    relevance: str  # 답변의 문서에 대한 관련성,
    is_question: bool  # 질문인지 확인

# 업스테이지 문서 관련성 체크 기능을 설정합니다. https://upstage.ai
upstage_ground_checker = UpstageGroundednessCheck()

# 문서에서 검색하여 관련성 있는 문서를 찾습니다.
def retrieve_document(state: GraphState) -> GraphState:
    # Question 에 대한 문서 검색을 retriever 로 수행합니다.
    retrieved_docs = retriever_tool.invoke(state["question"])
    # 검색된 문서를 context 키에 저장합니다.
    return GraphState(context=(retrieved_docs))


# Chain을 사용하여 답변을 생성합니다.
def llm_answer(state: GraphState) -> GraphState:
    reponse = chain.invoke(
            {"chat": state["question"], "Context": state["context"], "chat_history": state["chat_history"]}
        )
    question = reponse.is_question
    answer = reponse.answer

    return GraphState(
        answer=answer,
        context=state["context"],
        question=state["question"],
        chat_history=state["chat_history"],
        is_question=question
    )

# 질문인지 확인합니다.
def is_question(state: GraphState) -> GraphState:
    print("질문인지 확인" + str(state["is_question"]))
    if state["is_question"]:
        return True
    else:
        return False


    

# 관련성 체크를 실행합니다.
def relevance_check(state: GraphState) -> GraphState:
    # 관련성 체크를 실행합니다. 결과: grounded, notGrounded, notSure
    response = upstage_ground_checker.run(
        {"context": state["context"], "answer": state["answer"]}
    )
    
    print("답변\n" + state["answer"])
    print("context\n" + state["context"])
    print("관련성 체크 결과:" + response)
    
    return GraphState(
        relevance=response,
        context=state["context"],
        answer=state["answer"],
        question=state["question"],
        chat_history=state["chat_history"]
    )


# 관련성 체크 결과를 반환합니다.
def is_relevant(state: GraphState) -> GraphState:
    if state["relevance"] == "grounded":
        return "관련성 O"
    elif state["relevance"] == "notGrounded":
        return "관련성 X"
    elif state["relevance"] == "notSure":
        return "확인불가"

st.set_page_config(page_title="채팅 인터페이스", page_icon=":speech_balloon:")

@st.cache_resource
def setup_workflow():
    # langgraph.graph에서 StateGraph와 END를 가져옵니다.
    workflow = StateGraph(GraphState)

    # 노드들을 정의합니다.
    workflow.add_node("retrieve", retrieve_document)  # 에이전트 노드를 추가합니다.
    #workflow.add_node("is_question", is_question)  # 에이전트 노드를 추가합니다.
    workflow.add_node("llm_answer", llm_answer)  # 정보 검색 노드를 추가합니다.


    workflow.add_node(
        "relevance_check", relevance_check
    )  # 답변의 문서에 대한 관련성 체크 노드를 추가합니다.

    # 각 노드들을 연결합니다.
    workflow.add_edge("retrieve", "llm_answer")  # 검색 -> 질문 체크
   # workflow.add_edge("llm_answer", "relevance_check")  # 답변 -> 관련성 체크

    # 조건부 엣지를 추가합니다.
    workflow.add_conditional_edges(
        "llm_answer",  # 
        is_question,
        {
            False: END,  # 질문이 아니면 종료
            True: "relevance_check",  # 질문이면 관련성 검사로 진행
        },
    )

    # 조건부 엣지를 추가합니다.
    workflow.add_conditional_edges(
        "relevance_check",  # 관련성 체크 노드에서 나온 결과를 is_relevant 함수에 전달합니다.
        is_relevant,
        {
            "관련성 O": END,  # 관련성이 있으면 종료합니다.
            "관련성 X": "retrieve",  # 관련성이 없으면 다시 답변을 생성합니다.
            "확인불가": "retrieve",  # 관련성 체크 결과가 모호하다면 다시 답변을 생성합니다.
        },
    )

    # 시작점을 설정합니다.
    workflow.set_entry_point("retrieve")

    return workflow.compile(checkpointer=MemorySaver())

app = setup_workflow()


# #그래프 시각화
# try:
#     png_data = app.get_graph(xray=True).draw_mermaid_png()
    
#     # PNG 데이터 타입 확인
#     print(type(png_data))
    
#     # 데이터가 bytes 타입인지 확인
#     if isinstance(png_data, bytes):
#         with open("graph.png", "wb") as f:
#             f.write(png_data)
#         print("Graph saved as graph.png")
#     else:
#         print(f"Unexpected data type: {type(png_data)}")
#         # 데이터의 내용을 간단히 출력해 볼 수 있습니다
#         print(png_data[:100] if png_data else "Empty data")
# except Exception as e:
#     print(f"Graph visualization failed. Error: {e}")
#     import traceback
#     traceback.print_exc()


print("재실행")

# Streamlit 앱 설정


# 세션 상태 초기화
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

if 'retriever_tool' not in st.session_state:
    st.session_state.retriever_tool = initialize_vector_store()
retriever_tool = st.session_state.retriever_tool

# 채팅 인터페이스
st.title("하형운 대표와의 채팅")
st.write("메가커피 창업주인 하형운 대표님과의 채팅입니다.")

# 채팅 기록 표시
for message in st.session_state.chat_history:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])
        if message["role"] == "assistant" and "context" in message:
            with st.expander("참고한 내용"):
                st.write(message["context"])
                st.write("---")

# 사용자 입력
if prompt := st.chat_input("메시지를 입력하세요."):
    # 사용자 메시지를 즉시 표시
    st.session_state.chat_history.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.markdown(prompt)

    # 컨텍스트 검색
    context_docs = st.session_state.retriever_tool.invoke(prompt)

    # 채팅 기록을 문자열로 변환
    chat_history_str = "\n".join([f"{msg['role']}: {msg['content']}" for msg in st.session_state.chat_history])

    # recursion_limit: 최대 반복 횟수, thread_id: 실행 ID (구분용)
    config = RunnableConfig(recursion_limit=30, configurable={"thread_id": "SELF-RAG"})

    # GraphState 객체를 활용하여 질문을 입력합니다.
    inputs = GraphState(question=prompt)
    output = app.invoke(inputs, config=config)

    # 응답 생성

    # 응답 처리 및 표시
    processed_response = output["answer"].replace('"', '').replace('/', '').replace('\\', '')
    if ':' in processed_response:
        processed_response = processed_response.split(':', 1)[1].strip()
    
    
    final_response = claude_chain.invoke({"chat": prompt, "Context": context_docs, "response" : processed_response})

    final_processed_response = final_response.replace('"', '').replace('/', '').replace('\\', '')

    # 응답을 채팅 기록에 추가하고 표시
    st.session_state.chat_history.append({"role": "assistant", "content": final_processed_response, "context": context_docs})
    with st.chat_message("assistant"):
        st.markdown(processed_response)
        with st.expander("참고한 내용"):
            st.write(context_docs)
            st.write("---")

    # 페이지 새로고침
    st.rerun()